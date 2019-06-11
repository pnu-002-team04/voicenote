import io
import os
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from google.cloud import speech
from google.cloud.speech import enums
from google.cloud.speech import types
from datetime import datetime
import sys

var1 = sys.argv[1]
var2 = sys.argv[2]

now = datetime.now()
title = '%s-%s-%s Record File' % (now.year, now.month, now.day)
os.chdir(sys.argv[2])
class PDF(FPDF):
    def header(self):
        # Arial bold 15
        self.set_font('Arial', 'B', 15)
        # Calculate width of title and position
        w = self.get_string_width(title) + 6
        self.set_x((210 - w) / 2)
        # Background and text color
        self.set_draw_color(0, 80, 180)
        self.set_fill_color(230, 230, 0)
        self.set_text_color(220, 50, 50)
        # Thickness of frame (1 mm)
        self.set_line_width(1)
        # Title
        self.cell(w, 9, title, 1, 1, 'C', 1)
        # Line break
        self.ln(10)

    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        # Arial italic 8
        self.set_font('Arial', 'I', 8)
        # Text color in gray
        self.set_text_color(128)
        # Page number
        self.cell(0, 10, 'Page ' + str(self.page_no()), 0, 0, 'C')

    def chapter_title(self, num, label):
        # Arial 12
        self.set_font('Arial', '', 12)
        # Background color
        self.set_fill_color(200, 220, 255)
        # Title
        self.cell(0, 6, 'Chapter %d : %s' % (num, label), 0, 1, 'L', 1)
        # Line break
        self.ln(4)

    def chapter_body(self, name):
        # Read text file
        with open(name, 'rb') as fh:
            txt = fh.read().decode('latin-1')
        # Times 12
        self.set_font('Times', '', 12)
        # Output justified text
        self.multi_cell(0, 5, txt)
        # Line break
        self.ln()
        # Mention in italics
        self.set_font('', 'I')
        self.cell(0, 5, '(end of excerpt)')

    def print_chapter(self, num, title, name):
        self.add_page()
        self.chapter_title(num, title)
        self.chapter_body(name)

def transcribe_file():

        # # text to pdf
        pdf = PDF()
        pdf.set_title(title)
        pdf.set_author('PNU-TEAM-04')
        pdf.print_chapter(1, '', 'test.tmp')
        pdf.output('output.pdf', 'F')

        # text to doc
        document = Document()
        document.add_heading('%s-%s-%s Record File' % (now.year, now.month, now.day))
        document.add_paragraph(var1)
        document.save('output.docx')

if __name__ == '__main__':

    f = open("./test.tmp", 'w')
    f.write(var1)
    f.close()
    transcribe_file()
    print("EXIT") #trigger word

